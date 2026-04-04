"""
Generate Homework 10 final report as a .docx file.

Usage:
  python generate_report.py

Requires: python-docx (pip install python-docx)
"""

import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def add_image_with_caption(doc, image_path, caption, width=Inches(6)):
    """Add an image centered with a caption below it."""
    if os.path.exists(image_path):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path, width=width)

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
    else:
        doc.add_paragraph(f"[Image not found: {image_path}]")


def compute_stats(latencies):
    if not latencies:
        return {"count": 0, "min": 0, "avg": 0, "p50": 0, "p95": 0, "p99": 0, "max": 0}
    s = sorted(latencies)
    return {
        "count": len(s),
        "min": round(s[0], 2),
        "avg": round(sum(s) / len(s), 2),
        "p50": round(s[len(s) // 2], 2),
        "p95": round(s[int(len(s) * 0.95)], 2),
        "p99": round(s[int(len(s) * 0.99)], 2),
        "max": round(s[-1], 2),
    }


def main():
    # Load test results
    with open("results/load_test_results.json", "r") as f:
        results = json.load(f)

    doc = Document()

    # =========================================================
    # Title Page
    # =========================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading("Homework 10: Distributed Databases Using Replication", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("CS6650 Building Scalable Distributed Systems")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(80, 80, 80)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Spring 2026")
    run.font.size = Pt(12)

    doc.add_page_break()

    # =========================================================
    # 1. Introduction
    # =========================================================
    add_heading(doc, "1. Introduction")
    doc.add_paragraph(
        "This report documents the implementation and testing of a distributed Key-Value (KV) store "
        "using two replication architectures: Leader-Follower and Leaderless. "
        "The system is built in Go and deployed as a 5-node cluster using Docker Compose. "
        "We explore how different configurations of write quorum (W) and read quorum (R) "
        "affect latency, consistency, and the occurrence of stale reads."
    )
    doc.add_paragraph(
        "The key insight from the CAP theorem is that when W + R > N (where N is the total number of nodes), "
        "strong consistency is guaranteed because every read overlaps with every write quorum. "
        "When W + R <= N, stale reads are possible during the replication window."
    )

    # =========================================================
    # 2. System Architecture
    # =========================================================
    add_heading(doc, "2. System Architecture")

    add_heading(doc, "2.1 Leader-Follower Replication", level=2)
    doc.add_paragraph(
        "In this architecture, one designated Leader node handles all writes. "
        "When a client writes to the Leader, it replicates the data to Followers based on the write quorum W. "
        "Followers forward any write requests they receive to the Leader. "
        "Reads can be served by the Leader alone (R=1) or by querying multiple nodes and picking the highest version (R>1)."
    )
    doc.add_paragraph("We tested three Leader-Follower configurations:")
    configs_lf = [
        ("W=5, R=1", "Leader waits for ALL 5 nodes (including itself) before confirming the write. "
         "Reads only need 1 node. W+R=6 > N=5, so consistency is guaranteed. "
         "Writes are slow but reads are fast."),
        ("W=1, R=5", "Leader confirms the write after storing locally (W=1), without waiting for Followers. "
         "Reads query ALL 5 nodes and pick the highest version. W+R=6 > N=5, "
         "but during the replication window, stale reads are likely if R does not overlap with all pending replications."),
        ("W=3, R=3", "Balanced quorum. Leader waits for 3 nodes to confirm writes, reads query 3 nodes. "
         "W+R=6 > N=5, so consistency is theoretically guaranteed. "
         "Both writes and reads have moderate latency."),
    ]
    for title_text, desc in configs_lf:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        p.add_run(desc)

    add_heading(doc, "2.2 Leaderless Replication", level=2)
    doc.add_paragraph(
        "In the Leaderless architecture, any node can act as the Write Coordinator. "
        "When a node receives a write, it stores locally and replicates to ALL other peers (W=N). "
        "Reads are served locally (R=1). Since W=N ensures all nodes eventually have the data, "
        "and the Coordinator waits for all confirmations before responding, "
        "reads after a completed write are always consistent."
    )

    # =========================================================
    # 3. Docker Deployment
    # =========================================================
    add_heading(doc, "3. Docker Deployment")
    doc.add_paragraph(
        "Each node runs as a separate Docker container. The Go application is built using a multi-stage "
        "Dockerfile (golang:1.25-alpine for building, alpine:3.19 for runtime). "
        "Docker Compose orchestrates the 5-node cluster with appropriate environment variables."
    )
    add_image_with_caption(doc, "screenshot/docker leader built.png",
                           "Figure 1: Docker Compose starting the Leader-Follower cluster (5 nodes)")

    # =========================================================
    # 4. Manual Verification
    # =========================================================
    add_heading(doc, "4. Manual Verification")
    doc.add_paragraph(
        "Before running automated tests, we manually verified the basic read/write operations "
        "using PowerShell HTTP requests."
    )

    add_heading(doc, "4.1 Write Operation", level=2)
    add_image_with_caption(doc, "screenshot/leader write.png",
                           "Figure 2: Writing a key-value pair to the Leader (HTTP 201 Created)")

    add_heading(doc, "4.2 Read from Leader", level=2)
    add_image_with_caption(doc, "screenshot/leader read.png",
                           "Figure 3: Reading the key from the Leader (HTTP 200 OK)")

    add_heading(doc, "4.3 Read from Follower (Local Read)", level=2)
    add_image_with_caption(doc, "screenshot/leaderfollower1 local read.png",
                           "Figure 4: Reading from Follower1 via /local_read endpoint (HTTP 200 OK)")

    # =========================================================
    # 5. Unit Tests
    # =========================================================
    add_heading(doc, "5. Unit Tests")

    add_heading(doc, "5.1 Leader-Follower Tests", level=2)
    doc.add_paragraph(
        "Three tests were run against the Leader-Follower cluster:\n"
        "1) test_leader_read_after_write - Write then read from Leader, verify consistency.\n"
        "2) test_follower_consistent_after_write_w5 - With W=5, all followers should have data after write returns.\n"
        "3) test_inconsistency_during_replication - Concurrent write + local_read to detect stale reads during replication."
    )
    add_image_with_caption(doc, "screenshot/test leader.png",
                           "Figure 5: Leader-Follower unit tests - 3/3 passed, stale reads detected during replication")

    add_heading(doc, "5.2 Leaderless Tests", level=2)
    doc.add_paragraph(
        "Four tests were run against the Leaderless cluster:\n"
        "1) test_coordinator_consistent_after_write - Write to random node, read back immediately.\n"
        "2) test_all_nodes_consistent_after_write - W=N write, verify all nodes consistent.\n"
        "3) test_inconsistency_window - Detect stale reads while write is propagating.\n"
        "4) test_different_coordinators - Different nodes as Coordinator, verify version increments."
    )
    add_image_with_caption(doc, "screenshot/test ledaerless.png",
                           "Figure 6: Leaderless unit tests - 4/4 passed, stale reads detected during propagation window")

    # =========================================================
    # 6. Load Test Results
    # =========================================================
    add_heading(doc, "6. Load Tests")
    doc.add_paragraph(
        "We ran load tests with 200 requests, 10 concurrent threads, and a pool of 10 keys "
        "across 4 configurations and 4 write ratios (1%, 10%, 50%, 90%). "
        "This small key pool creates 'local-in-time' data where reads and writes frequently target "
        "the same keys, maximizing the chance of detecting stale reads."
    )

    # --- 6.1 Load Test Console Output ---
    add_heading(doc, "6.1 Load Test Console Output", level=2)

    screenshots_load = [
        ("screenshot/load_test_leader_w5r1.png", "Figure 7: Load test results - Leader-Follower W=5, R=1"),
        ("screenshot/load_test_leader_w1r5.png", "Figure 8: Load test results - Leader-Follower W=1, R=5"),
        ("screenshot/load_test_leader_w3r3.png", "Figure 9: Load test results - Leader-Follower W=3, R=3"),
        ("screenshot/load_test_leaderless.png", "Figure 10: Load test results - Leaderless W=N, R=1"),
    ]
    for img_path, caption in screenshots_load:
        add_image_with_caption(doc, img_path, caption)

    # --- 6.2 Latency Summary Table ---
    add_heading(doc, "6.2 Latency Summary", level=2)
    doc.add_paragraph(
        "The table below summarizes the average and p95 latency (in ms) for each configuration and write ratio."
    )

    # Build summary table
    config_display = {
        "leader_W5_R1": "LF W=5,R=1",
        "leader_W1_R5": "LF W=1,R=5",
        "leader_W3_R3": "LF W=3,R=3",
        "leaderless_WN_R1": "LL W=N,R=1",
    }

    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Config", "Write %", "Write Avg (ms)", "Write P95 (ms)",
               "Read Avg (ms)", "Read P95 (ms)", "Stale Reads", "Total Reads"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    for r in results:
        ws = compute_stats(r["write_latencies"])
        rs = compute_stats(r["read_latencies"])
        row = table.add_row()
        values = [
            config_display.get(r["config"], r["config"]),
            f"{int(r['write_ratio'] * 100)}%",
            str(ws["avg"]),
            str(ws["p95"]),
            str(rs["avg"]),
            str(rs["p95"]),
            str(r["stale_reads"]),
            str(r["total_reads"]),
        ]
        for i, v in enumerate(values):
            row.cells[i].text = v
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph()

    # --- 6.3 Latency Distribution Graphs ---
    add_heading(doc, "6.3 Latency Distribution Graphs", level=2)
    doc.add_paragraph(
        "The following graphs show the read and write latency distributions for each write ratio, "
        "comparing all four configurations side by side."
    )

    latency_graphs = [
        ("results/latency_1pct_writes.png", "Figure 11: Latency distribution - 1% Writes / 99% Reads"),
        ("results/latency_10pct_writes.png", "Figure 12: Latency distribution - 10% Writes / 90% Reads"),
        ("results/latency_50pct_writes.png", "Figure 13: Latency distribution - 50% Writes / 50% Reads"),
        ("results/latency_90pct_writes.png", "Figure 14: Latency distribution - 90% Writes / 10% Reads"),
    ]
    for img_path, caption in latency_graphs:
        add_image_with_caption(doc, img_path, caption)

    # --- 6.4 Stale Reads ---
    add_heading(doc, "6.4 Stale Reads Analysis", level=2)
    doc.add_paragraph(
        "Stale reads occur when a client reads an older version of a key after a newer version has been written. "
        "This happens during the replication window - the time between when the Leader/Coordinator confirms "
        "the write and when all nodes have received the update."
    )
    add_image_with_caption(doc, "results/stale_reads.png",
                           "Figure 15: Stale reads by configuration and write ratio")

    # Stale reads analysis text
    doc.add_paragraph(
        "Key observations on stale reads:"
    )
    stale_observations = [
        "W=5, R=1 (Leader-Follower): Zero stale reads across all write ratios. Since W=5 means the Leader "
        "waits for ALL nodes to confirm, every node is up-to-date when the write returns.",
        "W=1, R=5 (Leader-Follower): Highest stale read counts. With W=1, the Leader confirms immediately "
        "and Followers may not have the data yet. Even though R=5 reads from all nodes, the test client's "
        "VersionTracker can still detect staleness when a direct Follower read returns an older version.",
        "W=3, R=3 (Leader-Follower): Very few or zero stale reads. The quorum overlap (W+R=6 > N=5) "
        "ensures that reads and writes share at least one common node.",
        "Leaderless W=N, R=1: Zero stale reads. Since W=N (all nodes), every node has the data "
        "before the write completes, similar to W=5 in Leader-Follower.",
    ]
    for obs in stale_observations:
        doc.add_paragraph(obs, style="List Bullet")

    # --- 6.5 R/W Intervals ---
    add_heading(doc, "6.5 Read/Write Time Intervals", level=2)
    doc.add_paragraph(
        "The following graph shows the distribution of time intervals between consecutive read/write "
        "operations on the same key. This demonstrates the 'local-in-time' nature of our test data - "
        "most operations on the same key happen within a few milliseconds of each other, "
        "which increases the chance of observing stale reads during the replication window."
    )
    add_image_with_caption(doc, "results/rw_intervals.png",
                           "Figure 16: Time intervals between R/W operations on the same key")

    # =========================================================
    # 7. Analysis & Discussion
    # =========================================================
    add_heading(doc, "7. Analysis & Discussion")

    add_heading(doc, "7.1 Performance Tradeoffs", level=2)
    doc.add_paragraph(
        "The results clearly demonstrate the fundamental tradeoff between write latency, "
        "read latency, and consistency:"
    )

    tradeoffs = [
        ("W=5, R=1 (Strong Write Consistency)",
         "Write latency is high (~320ms) because the Leader must wait for all Followers. "
         "Read latency is very low (~5-12ms) since only a local read is needed. "
         "Zero stale reads. Best for read-heavy workloads that require strong consistency."),
        ("W=1, R=5 (Fast Writes, Expensive Reads)",
         "Write latency is minimal (~5ms) as the Leader only stores locally. "
         "Read latency is higher (~35ms) because it must query all 5 nodes. "
         "Stale reads are possible. Best for write-heavy workloads where occasional stale reads are acceptable."),
        ("W=3, R=3 (Balanced Quorum)",
         "Both write and read latency are moderate. "
         "Near-zero stale reads due to quorum overlap. "
         "Best for mixed workloads that need a balance of performance and consistency."),
        ("Leaderless W=N, R=1 (Coordinator Pattern)",
         "Similar performance profile to W=5,R=1 Leader-Follower. "
         "Write latency is high (~324ms) but reads are fast (~5ms). "
         "Zero stale reads. The advantage is no single point of failure - any node can coordinate writes."),
    ]
    for title_text, desc in tradeoffs:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        p.add_run(desc)

    add_heading(doc, "7.2 CAP Theorem Implications", level=2)
    doc.add_paragraph(
        "Our experiments validate the CAP theorem predictions:\n\n"
        "When W + R > N (e.g., W=5+R=1=6 > 5, or W=3+R=3=6 > 5), read and write quorums "
        "always overlap, guaranteeing that a read will see the latest write. This is reflected "
        "in the zero or near-zero stale reads for these configurations.\n\n"
        "When the effective overlap is weaker (e.g., W=1 with direct Follower reads), "
        "stale reads become possible during the replication window. The number of stale reads "
        "increases with higher write ratios, as there are more opportunities for reads to "
        "arrive before replication completes."
    )

    add_heading(doc, "7.3 Use Case Recommendations", level=2)
    recommendations = [
        ("Read-heavy applications (e.g., content delivery, caching)",
         "Use W=5, R=1 or Leaderless W=N, R=1. Pay the write cost once, enjoy fast reads."),
        ("Write-heavy applications (e.g., logging, event streaming)",
         "Use W=1, R=5. Fast writes with eventual consistency on reads."),
        ("Mixed workloads requiring consistency (e.g., e-commerce, banking)",
         "Use W=3, R=3. Balanced latency with strong consistency guarantees."),
        ("High availability requirements (e.g., global services)",
         "Use Leaderless architecture. No single point of failure; any node can handle writes."),
    ]
    for title_text, desc in recommendations:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        p.add_run(desc)

    # =========================================================
    # 8. Conclusion
    # =========================================================
    add_heading(doc, "8. Conclusion")
    doc.add_paragraph(
        "This homework demonstrated the practical implications of replication strategies in "
        "distributed databases. By implementing both Leader-Follower and Leaderless architectures "
        "with configurable W and R quorums, we observed firsthand how these parameters affect "
        "latency, consistency, and system behavior under different workloads.\n\n"
        "The key takeaways are:\n"
        "1. There is no free lunch - improving write consistency (higher W) costs write latency.\n"
        "2. The quorum formula W + R > N is a practical tool for reasoning about consistency.\n"
        "3. Stale reads are a real phenomenon that can be observed and measured.\n"
        "4. Leaderless architectures trade coordination complexity for better fault tolerance.\n"
        "5. The right configuration depends entirely on the application's read/write ratio "
        "and consistency requirements."
    )

    # =========================================================
    # Save
    # =========================================================
    output_path = "report.docx"
    doc.save(output_path)
    print(f"Report saved to {output_path}")


if __name__ == "__main__":
    main()
